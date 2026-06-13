from io import BytesIO

from docx import Document
from pypdf import PdfReader

from app.models.schemas import ParsedDocument, UploadedFileMeta
from app.utils.file_helpers import ALLOWED_EXTENSIONS, get_extension, hash_bytes


class FileParserService:
    """Handles upload validation and text extraction for documents."""

    def is_supported_filename(self, filename: str | None) -> bool:
        ext = get_extension(filename)
        return ext in ALLOWED_EXTENSIONS

    def parse_document(
        self,
        filename: str | None,
        content_type: str | None,
        file_bytes: bytes,
    ) -> ParsedDocument:
        safe_name = filename or "unknown"
        extension = get_extension(safe_name)
        extracted_text = self._extract_text(extension, file_bytes)
        if not extracted_text.strip():
            raise ValueError("No readable text found in file.")
        prepared_text = self._prepare_text(extracted_text)
        word_count = len(prepared_text.split())
        preview = prepared_text[:280]

        metadata = UploadedFileMeta(
            filename=safe_name,
            content_type=content_type or "application/octet-stream",
            extension=extension,
            size_bytes=len(file_bytes),
            file_hash=hash_bytes(file_bytes),
            parser_status="parsed",
            word_count=word_count,
            character_count=len(prepared_text),
            preview_snippet=preview,
        )

        return ParsedDocument(
            metadata=metadata,
            extracted_text=extracted_text,
            prepared_text=prepared_text,
        )

    def _extract_text(self, extension: str, file_bytes: bytes) -> str:
        if extension == "txt":
            return self._extract_txt(file_bytes)
        if extension == "docx":
            return self._extract_docx(file_bytes)
        if extension == "pdf":
            return self._extract_pdf(file_bytes)
        raise ValueError("Unsupported file type. Allowed: PDF, DOCX, TXT.")

    @staticmethod
    def _extract_txt(file_bytes: bytes) -> str:
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return file_bytes.decode("latin-1")
            except Exception as exc:
                raise ValueError("Failed to decode TXT file.") from exc

    @staticmethod
    def _extract_docx(file_bytes: bytes) -> str:
        try:
            doc = Document(BytesIO(file_bytes))
            lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            return "\n".join(lines)
        except Exception as exc:
            raise ValueError("Failed to parse DOCX file.") from exc

    @staticmethod
    def _extract_pdf(file_bytes: bytes) -> str:
        try:
            reader = PdfReader(BytesIO(file_bytes))
            pages = [(page.extract_text() or "").strip() for page in reader.pages]
            return "\n".join([p for p in pages if p])
        except Exception as exc:
            raise ValueError("Failed to parse PDF file.") from exc

    @staticmethod
    def _prepare_text(text: str) -> str:
        # Lightweight normalization so output is ready for downstream LLM analysis.
        normalized = " ".join(text.replace("\r", "\n").split())
        return normalized.strip()
